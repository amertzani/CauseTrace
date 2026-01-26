"""
DOCX Agent
==========

Specialized agent for processing DOCX/DOC files.
Uses python-docx for text extraction.
"""

import os
from docx import Document
from typing import Dict
from .base_agent import BaseAgent


class DOCXAgent(BaseAgent):
    """Agent specialized in processing DOCX/DOC files."""
    
    def __init__(self):
        super().__init__("DOCX Agent")
        self.supported_extensions = ['.docx', '.DOCX', '.doc', '.DOC']
    
    def can_process(self, file_path: str) -> bool:
        """Check if file is a DOCX or DOC."""
        ext = os.path.splitext(file_path)[1].lower()
        return ext in ['.docx', '.doc']
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts).strip()
            
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    def extract_metadata(self, file_path: str) -> Dict:
        """Extract DOCX-specific metadata."""
        metadata = super().extract_metadata(file_path)
        
        try:
            doc = Document(file_path)
            metadata['paragraphs'] = len([p for p in doc.paragraphs if p.text.strip()])
            metadata['tables'] = len(doc.tables)
            
            # Try to get document properties
            if doc.core_properties:
                metadata['doc_title'] = doc.core_properties.title or ''
                metadata['doc_author'] = doc.core_properties.author or ''
                metadata['doc_subject'] = doc.core_properties.subject or ''
        except Exception:
            # If metadata extraction fails, continue without it
            metadata['paragraphs'] = 0
            metadata['tables'] = 0
        
        return metadata

